"""Unit tests for per-format document loaders."""

from __future__ import annotations

import docx
import pytest
from fpdf import FPDF

from app.core.exceptions import IngestionError, UnsupportedFileTypeError
from app.rag.loaders.html import HtmlLoader
from app.rag.loaders.markdown import MarkdownLoader
from app.rag.loaders.pdf import PdfLoader
from app.rag.loaders.registry import get_loader
from app.rag.loaders.text import TextLoader


class TestTextLoader:
    def test_loads_content(self, tmp_path) -> None:
        path = tmp_path / "note.txt"
        path.write_text("Employees may expense meals up to $50 per day.")
        elements = TextLoader().load(path)
        assert len(elements) == 1
        assert "expense meals" in elements[0].content

    def test_empty_file_raises(self, tmp_path) -> None:
        path = tmp_path / "empty.txt"
        path.write_text("   ")
        with pytest.raises(IngestionError):
            TextLoader().load(path)


class TestMarkdownLoader:
    def test_sections_from_headings(self, tmp_path) -> None:
        path = tmp_path / "policy.md"
        path.write_text(
            "# Vacation\nEmployees get 20 days.\n\n"
            "## Sick Leave\nUnlimited with a doctor's note.\n"
        )
        elements = MarkdownLoader().load(path)
        assert [e.section for e in elements] == ["Vacation", "Sick Leave"]
        assert "20 days" in elements[0].content

    def test_headings_inside_code_blocks_ignored(self, tmp_path) -> None:
        path = tmp_path / "code.md"
        path.write_text("# Real\ntext\n```\n# not a heading\n```\nmore\n")
        elements = MarkdownLoader().load(path)
        assert len(elements) == 1
        assert elements[0].section == "Real"
        assert "# not a heading" in elements[0].content


class TestHtmlLoader:
    def test_sections_and_noise_removal(self, tmp_path) -> None:
        path = tmp_path / "page.html"
        path.write_text(
            "<html><head><style>body{}</style></head><body>"
            "<script>alert(1)</script>"
            "<h1>Benefits</h1><p>Health insurance is covered at 100%.</p>"
            "<h2>Dental</h2><p>Dental is covered at 80%.</p>"
            "</body></html>"
        )
        elements = HtmlLoader().load(path)
        text = " ".join(e.content for e in elements)
        assert "alert(1)" not in text
        assert "Health insurance" in text
        sections = [e.section for e in elements]
        assert "Benefits" in sections
        assert "Dental" in sections

    def test_unstructured_html_falls_back_to_raw_text(self, tmp_path) -> None:
        path = tmp_path / "raw.html"
        path.write_text("<html><body><div>Just a div with policy text.</div></body></html>")
        elements = HtmlLoader().load(path)
        assert any("policy text" in e.content for e in elements)


class TestDocxLoader:
    def test_headings_become_sections(self, tmp_path) -> None:
        path = tmp_path / "doc.docx"
        document = docx.Document()
        document.add_heading("Remote Work", level=1)
        document.add_paragraph("Up to three days per week remote.")
        document.add_heading("Equipment", level=1)
        document.add_paragraph("Laptops are provided by IT.")
        document.save(str(path))

        from app.rag.loaders.docx import DocxLoader

        elements = DocxLoader().load(path)
        assert [e.section for e in elements] == ["Remote Work", "Equipment"]
        assert "three days" in elements[0].content


class TestPdfLoader:
    def test_pages_have_numbers(self, tmp_path) -> None:
        path = tmp_path / "doc.pdf"
        pdf = FPDF()
        pdf.set_font("helvetica", size=12)
        pdf.add_page()
        pdf.multi_cell(0, 10, "Page one: travel expenses must be pre-approved.")
        pdf.add_page()
        pdf.multi_cell(0, 10, "Page two: use the corporate card for bookings.")
        pdf.output(str(path))

        elements = PdfLoader().load(path)
        assert [e.page for e in elements] == [1, 2]
        assert "pre-approved" in elements[0].content

    def test_garbage_bytes_raise_ingestion_error(self, tmp_path) -> None:
        path = tmp_path / "broken.pdf"
        path.write_bytes(b"this is not a pdf at all")
        with pytest.raises(IngestionError):
            PdfLoader().load(path)


class TestRegistry:
    def test_known_extensions(self) -> None:
        assert get_loader(".pdf").__class__.__name__ == "PdfLoader"
        assert get_loader(".MD").__class__.__name__ == "MarkdownLoader"
        assert get_loader(".htm").__class__.__name__ == "HtmlLoader"

    def test_unknown_extension_raises(self) -> None:
        with pytest.raises(UnsupportedFileTypeError):
            get_loader(".exe")
